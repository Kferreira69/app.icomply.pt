"""
Generate the iComply Academy Video Scripts Word document (20 videos)
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = r"C:\Users\cmmfe\OneDrive\2022 Contemporary Constellation\8_iComply\NEW_iComply\RFP\iComply_Unicorn_SeriesA_Documentation_Rules\iComply_Academy_Video_Scripts.docx"

BRAND_BLUE   = RGBColor(0x1E, 0x3A, 0x8A)
BRAND_LIGHT  = RGBColor(0x3B, 0x82, 0xF6)
DARK_TEXT    = RGBColor(0x11, 0x18, 0x27)
GRAY_TEXT    = RGBColor(0x6B, 0x72, 0x80)
GREEN_TEXT   = RGBColor(0x16, 0xA3, 0x4A)
WHITE        = RGBColor(0xFF, 0xFF, 0xFF)

VIDEOS = [
    {
        "num": 1,
        "title": "Bem-vindo ao iComply",
        "category": "Introdução",
        "level": "Iniciante",
        "duration": "3 min",
        "objective": "Apresentar a plataforma iComply, os seus módulos principais e a filosofia de GRC (Governance, Risk & Compliance) integrado.",
        "topics": ["O que é GRC", "Visão geral dos módulos", "Navegação inicial"],
        "script": [
            ("Cena 1 — Ecrã de boas-vindas (0:00–0:30)", [
                "Narrador: Bem-vindo ao iComply — a plataforma de Governance, Risk & Compliance desenhada para empresas europeias.",
                "Ação: Mostrar o ecrã de login. O utilizador entra com as suas credenciais e chega ao dashboard.",
                "Narrador: Com o iComply pode gerir riscos, documentos, auditorias, conformidade com RGPD, NIS2 e ISO 27001 — tudo num só lugar.",
            ]),
            ("Cena 2 — Dashboard principal (0:30–1:30)", [
                "Ação: Percorrer os widgets do dashboard: score de conformidade, riscos críticos, tarefas pendentes, alertas regulatórios.",
                "Narrador: O dashboard dá-lhe uma visão executiva em tempo real. O score de conformidade reflete o estado actual de todos os frameworks activos.",
                "Narrador: Os widgets são configuráveis — pode reorganizá-los conforme as suas prioridades.",
            ]),
            ("Cena 3 — Menu lateral (1:30–2:30)", [
                "Ação: Expandir o menu lateral, mostrar as secções: Riscos, Documentos, Auditorias, RGPD, Tarefas, iGuard, etc.",
                "Narrador: O menu lateral dá acesso a todos os módulos. Vamos explorar cada um nos vídeos seguintes.",
            ]),
            ("Cena 4 — Resumo e próximos passos (2:30–3:00)", [
                "Narrador: Nos próximos vídeos veremos como registar e gerir riscos, criar documentos de política, lançar auditorias e muito mais.",
                "Ação: Mostrar o Centro de Formação com os vídeos disponíveis.",
            ]),
        ],
    },
    {
        "num": 2,
        "title": "Registar e Gerir Riscos",
        "category": "Gestão de Riscos",
        "level": "Iniciante",
        "duration": "5 min",
        "objective": "Demonstrar como criar, avaliar e gerir riscos no módulo de gestão de riscos do iComply.",
        "topics": ["Criar um risco", "Matriz de risco (probabilidade × impacto)", "Atribuir responsável", "Risco residual vs. inerente"],
        "script": [
            ("Cena 1 — Aceder ao módulo de Riscos (0:00–0:30)", [
                "Narrador: Comece por clicar em 'Riscos' no menu lateral.",
                "Ação: Mostrar a listagem de riscos com filtros e matriz visual.",
            ]),
            ("Cena 2 — Criar um novo risco (0:30–2:30)", [
                "Ação: Clicar em 'Novo Risco'. Preencher: Título 'Acesso não autorizado a dados de clientes', Categoria 'Segurança de Informação', Descrição detalhada.",
                "Narrador: Dê ao risco um título claro e uma descrição que explique a natureza da ameaça.",
                "Ação: Definir Probabilidade: 3 (Provável), Impacto: 4 (Severo). O sistema calcula automaticamente o score: 12 — Risco Alto.",
                "Narrador: A matriz de risco combina probabilidade e impacto para classificar o risco como Baixo, Médio, Alto ou Crítico.",
                "Ação: Selecionar o framework aplicável (ex: ISO 27001, RGPD). Atribuir responsável. Definir data de revisão.",
            ]),
            ("Cena 3 — Tratamento e controlos (2:30–4:00)", [
                "Narrador: Para cada risco pode definir a estratégia de tratamento: Mitigar, Aceitar, Transferir ou Evitar.",
                "Ação: Selecionar 'Mitigar'. Adicionar controlo: 'Implementar autenticação multi-factor'. Definir prazo e responsável.",
                "Narrador: Após implementar os controlos, o risco residual deve ser reavaliado.",
                "Ação: Actualizar probabilidade residual para 1 (Raro). Score residual: 4 — Risco Baixo.",
            ]),
            ("Cena 4 — Dashboard de riscos (4:00–5:00)", [
                "Narrador: No dashboard de riscos pode ver a distribuição por categoria, por framework e os riscos que requerem atenção imediata.",
                "Ação: Mostrar gráficos e filtros da listagem. Exportar relatório em PDF.",
            ]),
        ],
    },
    {
        "num": 3,
        "title": "Gestão de Documentos e Políticas",
        "category": "Documentação",
        "level": "Iniciante",
        "duration": "4 min",
        "objective": "Aprender a criar, versionar e publicar documentos de política e procedimento no iComply.",
        "topics": ["Criar documento", "Controlo de versões", "Aprovação e publicação", "Revisões periódicas"],
        "script": [
            ("Cena 1 — Biblioteca de documentos (0:00–0:45)", [
                "Narrador: Aceda a 'Documentos' no menu lateral. Aqui encontra todas as políticas, procedimentos e templates da sua organização.",
                "Ação: Mostrar a listagem com filtros por tipo, estado (Rascunho/Aprovado/Publicado) e data.",
            ]),
            ("Cena 2 — Criar novo documento (0:45–2:00)", [
                "Ação: Clicar 'Novo Documento'. Selecionar tipo: 'Política de Segurança da Informação'. Preencher título e descrição.",
                "Narrador: O editor suporta formatação rica. Pode também carregar ficheiros existentes em Word ou PDF.",
                "Ação: Escrever conteúdo. Associar ao framework ISO 27001. Definir proprietário do documento.",
            ]),
            ("Cena 3 — Fluxo de aprovação (2:00–3:15)", [
                "Narrador: Quando o documento estiver pronto, submeta-o para aprovação.",
                "Ação: Clicar 'Enviar para aprovação'. Selecionar aprovadores. Eles recebem notificação.",
                "Narrador: Após aprovação, o documento é automaticamente publicado e o número de versão incrementado.",
                "Ação: Mostrar histórico de versões do documento.",
            ]),
            ("Cena 4 — Revisões periódicas (3:15–4:00)", [
                "Narrador: O iComply alerta automaticamente quando um documento se aproxima da data de revisão.",
                "Ação: Mostrar alerta de revisão no dashboard e como actualizar o documento.",
            ]),
        ],
    },
    {
        "num": 4,
        "title": "Planeamento e Execução de Auditorias",
        "category": "Auditorias",
        "level": "Intermédio",
        "duration": "6 min",
        "objective": "Criar e gerir programas de auditoria interna, atribuir tarefas e registar evidências.",
        "topics": ["Criar programa de auditoria", "Atribuir auditores", "Recolha de evidências", "Relatório de auditoria"],
        "script": [
            ("Cena 1 — Módulo de Auditorias (0:00–0:45)", [
                "Narrador: As auditorias internas são essenciais para verificar a eficácia dos controlos. Aceda ao módulo 'Auditorias'.",
                "Ação: Mostrar o calendário de auditorias e as auditorias em curso.",
            ]),
            ("Cena 2 — Criar nova auditoria (0:45–2:30)", [
                "Ação: Clicar 'Nova Auditoria'. Dar nome: 'Auditoria ISO 27001 — Q2 2025'. Selecionar scope e framework.",
                "Narrador: Defina o âmbito, os critérios de auditoria e o período.",
                "Ação: Selecionar o template de checklist ISO 27001. O sistema carrega automaticamente os controlos relevantes.",
                "Narrador: Cada controlo será avaliado durante a auditoria: Conforme, Não Conforme ou Não Aplicável.",
            ]),
            ("Cena 3 — Atribuição e execução (2:30–4:15)", [
                "Narrador: Atribua a auditoria a um auditor interno ou convide um auditor externo.",
                "Ação: Selecionar auditor. O auditor recebe notificação e acesso ao portal.",
                "Ação: Mostrar a vista do auditor — avaliando cada controlo, adicionando notas e carregando evidências.",
                "Narrador: O sistema regista todas as evidências com metadata: data, responsável e ficheiro.",
            ]),
            ("Cena 4 — Não conformidades e plano de acção (4:15–5:15)", [
                "Narrador: Quando um controlo é marcado como Não Conforme, deve criar uma não conformidade e definir acções correctivas.",
                "Ação: Criar NC, atribuir responsável, prazo e acção correctiva.",
            ]),
            ("Cena 5 — Relatório final (5:15–6:00)", [
                "Narrador: No final da auditoria, o iComply gera automaticamente um relatório completo em PDF.",
                "Ação: Mostrar o relatório de auditoria gerado, com sumário executivo, resultados e acções correctivas.",
            ]),
        ],
    },
    {
        "num": 5,
        "title": "RGPD — Registo de Actividades de Tratamento",
        "category": "RGPD",
        "level": "Intermédio",
        "duration": "5 min",
        "objective": "Criar e manter o Registo das Actividades de Tratamento (RAT) conforme exigido pelo RGPD.",
        "topics": ["Criar actividade de tratamento", "Base legal", "Transferências internacionais", "Exportar RAT"],
        "script": [
            ("Cena 1 — Módulo RGPD (0:00–0:30)", [
                "Narrador: O módulo RGPD do iComply ajuda-o a cumprir o Regulamento Geral sobre a Protecção de Dados.",
                "Ação: Aceder ao módulo RGPD, mostrar as secções: RAT, Pedidos de Titulares, DPIAs, Violações.",
            ]),
            ("Cena 2 — Criar actividade de tratamento (0:30–2:30)", [
                "Ação: Clicar 'Nova Actividade'. Preencher: Nome 'Gestão de Recursos Humanos', Responsável pelo tratamento, Categorias de dados.",
                "Narrador: Especifique as categorias de dados pessoais tratados — por exemplo: identificação, dados profissionais, dados financeiros.",
                "Ação: Definir finalidade, base legal (Execução de contrato), período de retenção (5 anos), destinatários.",
                "Narrador: A base legal é fundamental. Escolha entre as 6 bases previstas no RGPD: consentimento, contrato, obrigação legal, etc.",
            ]),
            ("Cena 3 — Medidas de segurança e transferências (2:30–4:00)", [
                "Narrador: Para cada actividade, documente as medidas de segurança implementadas.",
                "Ação: Adicionar medidas: encriptação em repouso, controlo de acessos, pseudonimização.",
                "Narrador: Se transfere dados fora da UE/EEE, deve indicar as salvaguardas adequadas.",
            ]),
            ("Cena 4 — Exportar o RAT (4:00–5:00)", [
                "Narrador: O iComply gera automaticamente o RAT em formato Word, pronto para entrega à autoridade de supervisão.",
                "Ação: Clicar 'Exportar RAT'. Mostrar o documento gerado.",
            ]),
        ],
    },
    {
        "num": 6,
        "title": "Gestão de Tarefas e Fluxos de Aprovação",
        "category": "Tarefas",
        "level": "Iniciante",
        "duration": "4 min",
        "objective": "Criar, atribuir e acompanhar tarefas de conformidade com fluxos de aprovação configuráveis.",
        "topics": ["Criar tarefa", "Atribuir e priorizar", "Aprovações e revisões", "Dashboard de tarefas"],
        "script": [
            ("Cena 1 — Módulo de Tarefas (0:00–0:30)", [
                "Narrador: O módulo de Tarefas centraliza todo o trabalho de conformidade: revisões, implementações, seguimentos.",
                "Ação: Mostrar o board Kanban e a vista de lista.",
            ]),
            ("Cena 2 — Criar e configurar tarefa (0:30–2:00)", [
                "Ação: Clicar 'Nova Tarefa'. Preencher: Título, Descrição, Prioridade (Crítica/Alta/Média/Baixa), Data de vencimento.",
                "Narrador: Pode associar a tarefa a um risco, controlo ou documento específico.",
                "Ação: Atribuir a um colega. Adicionar checklist de sub-tarefas.",
            ]),
            ("Cena 3 — Aprovações (2:00–3:15)", [
                "Narrador: Algumas tarefas requerem aprovação antes de serem marcadas como concluídas.",
                "Ação: Activar 'Requer aprovação'. Selecionar aprovador. Quando a tarefa é submetida, o aprovador recebe notificação.",
                "Narrador: O aprovador pode aceitar ou devolver com comentários.",
            ]),
            ("Cena 4 — Dashboard e relatórios (3:15–4:00)", [
                "Narrador: O dashboard de tarefas mostra o progresso, tarefas em atraso e carga de trabalho por colaborador.",
                "Ação: Mostrar gráficos de progresso e filtros por responsável, data, prioridade.",
            ]),
        ],
    },
    {
        "num": 7,
        "title": "Recolha e Gestão de Evidências",
        "category": "Evidências",
        "level": "Intermédio",
        "duration": "4 min",
        "objective": "Demonstrar como recolher, organizar e vincular evidências a controlos e auditorias.",
        "topics": ["Carregar evidências manualmente", "Recolha automática (integrações)", "Vincular a controlos", "Exportar para auditoria"],
        "script": [
            ("Cena 1 — O papel das evidências (0:00–0:30)", [
                "Narrador: As evidências são a prova de que os controlos estão implementados e funcionam. O iComply centraliza todas as evidências.",
            ]),
            ("Cena 2 — Carregar evidência manual (0:30–1:45)", [
                "Ação: Na secção Evidências, clicar 'Nova Evidência'. Carregar ficheiro (captura de ecrã, log, certificado).",
                "Narrador: Adicione título, descrição, data de recolha e o controlo ou framework a que se aplica.",
            ]),
            ("Cena 3 — Recolha automática (1:45–3:00)", [
                "Narrador: Com as integrações activas, o iComply recolhe evidências automaticamente de ferramentas como Microsoft 365, AWS, GitHub.",
                "Ação: Mostrar a secção de integrações e as evidências recolhidas automaticamente com timestamp.",
                "Narrador: Cada evidência automática inclui metadata de origem, garantindo rastreabilidade total.",
            ]),
            ("Cena 4 — Exportar para auditoria (3:00–4:00)", [
                "Narrador: Durante uma auditoria, pode vincular evidências directamente a cada controlo avaliado.",
                "Ação: Na auditoria, seleccionar um controlo, 'Adicionar evidência', seleccionar da biblioteca.",
            ]),
        ],
    },
    {
        "num": 8,
        "title": "NIS2 — Implementação e Monitorização",
        "category": "NIS2",
        "level": "Avançado",
        "duration": "7 min",
        "objective": "Guiar na implementação do framework NIS2 (Directiva sobre Segurança das Redes e Sistemas de Informação).",
        "topics": ["Âmbito e obrigações NIS2", "Mapeamento de controlos", "Gestão de incidentes", "Relatório de conformidade NIS2"],
        "script": [
            ("Cena 1 — Introdução ao NIS2 (0:00–1:00)", [
                "Narrador: A Directiva NIS2 (EU 2022/2555) impõe requisitos de cibersegurança a operadores de serviços essenciais e importantes em Portugal e na UE.",
                "Narrador: O iComply inclui um framework NIS2 pré-configurado com todos os controlos relevantes.",
            ]),
            ("Cena 2 — Activar o framework NIS2 (1:00–2:30)", [
                "Ação: Em Frameworks, selecionar 'NIS2'. Clicar 'Activar'. O sistema cria automaticamente os controlos NIS2 na organização.",
                "Narrador: Os controlos NIS2 cobrem 10 domínios: gestão de riscos, resposta a incidentes, segurança da cadeia de fornecimento, etc.",
            ]),
            ("Cena 3 — Avaliar conformidade (2:30–4:30)", [
                "Ação: Percorrer os controlos NIS2. Para cada controlo: avaliar estado (Implementado/Parcial/Não Implementado).",
                "Narrador: Para cada controlo não implementado, crie uma tarefa com responsável e prazo.",
                "Ação: Vincular políticas e evidências existentes aos controlos NIS2 relevantes.",
            ]),
            ("Cena 4 — Gestão de incidentes NIS2 (4:30–6:00)", [
                "Narrador: A NIS2 obriga à notificação de incidentes significativos às autoridades competentes em 24 horas.",
                "Ação: Mostrar o registo de incidentes. Criar incidente, preencher campos obrigatórios NIS2: impacto, causa, acções tomadas.",
                "Narrador: O iComply gera automaticamente o relatório de notificação no formato exigido.",
            ]),
            ("Cena 5 — Relatório de conformidade (6:00–7:00)", [
                "Narrador: O relatório de conformidade NIS2 mostra o seu nível de maturidade em cada domínio.",
                "Ação: Gerar relatório NIS2. Mostrar o score por domínio e o roadmap de melhorias.",
            ]),
        ],
    },
    {
        "num": 9,
        "title": "ISO 27001 — Gestão da Segurança da Informação",
        "category": "ISO 27001",
        "level": "Avançado",
        "duration": "8 min",
        "objective": "Implementar o SGSI (Sistema de Gestão da Segurança da Informação) segundo a ISO 27001:2022.",
        "topics": ["Estrutura ISO 27001:2022", "Declaração de Aplicabilidade (SoA)", "Gestão de riscos ISO 27001", "Preparação para certificação"],
        "script": [
            ("Cena 1 — ISO 27001 no iComply (0:00–1:00)", [
                "Narrador: A ISO 27001 é o padrão internacional para gestão da segurança da informação. O iComply inclui todos os 93 controlos do Anexo A da versão 2022.",
            ]),
            ("Cena 2 — Declaração de Aplicabilidade (1:00–3:30)", [
                "Narrador: A SoA (Statement of Applicability) é um documento central da ISO 27001 que justifica quais controlos se aplicam à sua organização.",
                "Ação: Aceder à SoA. Para cada controlo dos 93, marcar como: Aplicável / Não Aplicável + justificação.",
                "Ação: Para controlos aplicáveis, indicar se estão Implementados / Em curso / Planeados.",
                "Narrador: O iComply gera automaticamente a SoA em Word, pronta para auditoria de certificação.",
            ]),
            ("Cena 3 — Gestão de riscos ISO 27001 (3:30–5:30)", [
                "Narrador: A ISO 27001 exige uma metodologia formal de gestão de riscos. O iComply implementa a abordagem risk-based conforme a cláusula 6.1.",
                "Ação: Mostrar os riscos associados a activos de informação. Mapear riscos a controlos do Anexo A.",
            ]),
            ("Cena 4 — Plano de tratamento de riscos (5:30–7:00)", [
                "Ação: Gerar o Plano de Tratamento de Riscos. Mostrar o documento com todos os riscos, controlos associados e responsáveis.",
            ]),
            ("Cena 5 — Preparação para auditoria (7:00–8:00)", [
                "Narrador: Quando estiver pronto para a certificação, use o módulo de Auditorias para simular uma auditoria interna ISO 27001.",
                "Ação: Criar auditoria interna ISO 27001. Mostrar o relatório de pré-certificação.",
            ]),
        ],
    },
    {
        "num": 10,
        "title": "Dashboard Executivo e Board Reports",
        "category": "Dashboard",
        "level": "Intermédio",
        "duration": "4 min",
        "objective": "Usar o dashboard executivo e gerar relatórios para a Administração e o Conselho de Administração.",
        "topics": ["Widgets configuráveis", "KPIs de conformidade", "Relatório Board", "Exportação e impressão"],
        "script": [
            ("Cena 1 — Dashboard configurável (0:00–1:30)", [
                "Narrador: O dashboard executivo do iComply é totalmente configurável. Pode adicionar, remover e reorganizar widgets.",
                "Ação: Mostrar o modo de edição do dashboard. Adicionar widget 'Score de Conformidade', 'Riscos por Framework', 'Tarefas em Atraso'.",
            ]),
            ("Cena 2 — KPIs e tendências (1:30–2:30)", [
                "Narrador: Os KPIs mostram a evolução da conformidade ao longo do tempo — fundamental para demonstrar progresso à Administração.",
                "Ação: Mostrar gráfico de tendência do score de conformidade nos últimos 12 meses.",
            ]),
            ("Cena 3 — Board Report (2:30–3:30)", [
                "Narrador: O módulo Board Reports gera um relatório executivo profissional para apresentar ao Conselho de Administração.",
                "Ação: Aceder a Board Reports. Selecionar período. O sistema gera automaticamente o relatório com gráficos, riscos críticos e plano de acção.",
                "Narrador: O relatório pode ser exportado em PDF ou impresso directamente do browser.",
            ]),
            ("Cena 4 — Partilha e colaboração (3:30–4:00)", [
                "Narrador: Pode partilhar o dashboard com utilizadores externos, como auditores ou membros do Conselho, através de um link seguro.",
            ]),
        ],
    },
    {
        "num": 11,
        "title": "Motor de Automações",
        "category": "Automações",
        "level": "Avançado",
        "duration": "5 min",
        "objective": "Configurar automações para tarefas repetitivas: alertas, revisões periódicas e notificações.",
        "topics": ["Criar regra de automação", "Triggers e acções", "Automação de revisões", "Auditoria de automações"],
        "script": [
            ("Cena 1 — Motor de automações (0:00–0:45)", [
                "Narrador: As automações do iComply reduzem o trabalho manual ao automatizar tarefas recorrentes de conformidade.",
                "Ação: Aceder a Automações. Mostrar as automações activas e os seus logs.",
            ]),
            ("Cena 2 — Criar automação (0:45–2:30)", [
                "Ação: Clicar 'Nova Automação'. Nomear: 'Alerta de Documentos a Expirar'.",
                "Narrador: Defina o trigger: quando a data de revisão de um documento for em menos de 30 dias.",
                "Ação: Selecionar trigger 'Documento — data de revisão próxima'. Definir condição: 30 dias antes.",
                "Narrador: A acção: criar uma tarefa de revisão atribuída ao proprietário do documento.",
                "Ação: Definir acção 'Criar tarefa'. Configurar destinatário, prioridade, título da tarefa.",
            ]),
            ("Cena 3 — Templates de automação (2:30–3:45)", [
                "Narrador: O iComply inclui templates pré-configurados para os casos de uso mais comuns.",
                "Ação: Mostrar templates: Revisão trimestral de riscos, Relatório mensal de conformidade, Alerta de tarefas em atraso.",
            ]),
            ("Cena 4 — Monitorizar automações (3:45–5:00)", [
                "Narrador: O log de execução mostra todas as automações disparadas, os resultados e eventuais erros.",
                "Ação: Mostrar o log de execução de uma automação. Filtrar por data e estado.",
            ]),
        ],
    },
    {
        "num": 12,
        "title": "Centro de Formação — Como Usar",
        "category": "Centro de Formação",
        "level": "Iniciante",
        "duration": "3 min",
        "objective": "Apresentar o Centro de Formação do iComply e como utilizá-lo para formação contínua da equipa.",
        "topics": ["Navegar no catálogo", "Filtrar por nível e categoria", "Marcar progresso", "Atribuir formação à equipa"],
        "script": [
            ("Cena 1 — Aceder ao Centro de Formação (0:00–0:30)", [
                "Narrador: O Centro de Formação do iComply disponibiliza vídeos curtos e práticos sobre todos os módulos da plataforma.",
                "Ação: Clicar em 'Centro de Formação' no menu lateral.",
            ]),
            ("Cena 2 — Navegar no catálogo (0:30–1:30)", [
                "Ação: Mostrar os 20 vídeos organizados por categoria. Usar a pesquisa: escrever 'RGPD' e ver os vídeos filtrados.",
                "Narrador: Pode filtrar por nível — Iniciante, Intermédio ou Avançado — e por categoria de módulo.",
            ]),
            ("Cena 3 — Assistir e progredir (1:30–2:30)", [
                "Ação: Clicar num vídeo. Mostrar o player. Marcar como 'Visto'.",
                "Narrador: O sistema regista o seu progresso. Pode retomar de onde ficou ou rever os vídeos quantas vezes quiser.",
            ]),
            ("Cena 4 — Formação de equipa (2:30–3:00)", [
                "Narrador: Como administrador, pode atribuir vídeos de formação obrigatórios à sua equipa e acompanhar o progresso.",
                "Ação: Mostrar a função de atribuição de formação no perfil de utilizador.",
            ]),
        ],
    },
    {
        "num": 13,
        "title": "iGuard — Conformidade de Endpoints",
        "category": "iGuard",
        "level": "Intermédio",
        "duration": "5 min",
        "objective": "Instalar e configurar o agente iGuard nos dispositivos dos colaboradores para monitorizar conformidade.",
        "topics": ["Instalar o agente", "Gerar token", "Ver resultados no dashboard", "Revogar dispositivo"],
        "script": [
            ("Cena 1 — O que é o iGuard (0:00–0:45)", [
                "Narrador: O iGuard é um agente leve que corre nos dispositivos dos colaboradores e verifica automaticamente o estado de segurança: encriptação, screen lock, antivírus e actualizações.",
                "Narrador: Nunca acede a ficheiros pessoais, histórico de navegação ou comunicações.",
            ]),
            ("Cena 2 — Instalar o agente (0:45–2:30)", [
                "Ação: Aceder a iGuard → Instalar. Gerar token. Copiar o token.",
                "Narrador: O token é usado para associar o dispositivo à sua conta. Guarde-o — é mostrado apenas uma vez.",
                "Ação: Mostrar o tab Windows: descarregar iguard.exe, executar como administrador com o parâmetro /?token=TOKEN.",
                "Ação: Mostrar o tab macOS: executar o curl one-liner no Terminal.",
            ]),
            ("Cena 3 — Dashboard de dispositivos (2:30–4:00)", [
                "Narrador: Após a primeira verificação, o dispositivo aparece no dashboard iGuard com o seu score de conformidade.",
                "Ação: Mostrar o dashboard com os dispositivos da organização. Score ring, badges de verificação.",
                "Narrador: Pode ver quais verificações falharam e criar tarefas de remediação directamente.",
            ]),
            ("Cena 4 — Servidores e rede (4:00–5:00)", [
                "Narrador: Para além de endpoints, o iGuard também monitoriza servidores Linux e Windows Server, verificando configuração SSH, firewall e patches.",
                "Ação: Mostrar o tab 'Servidores' no dashboard e o tab 'Rede' com as Network Probes.",
            ]),
        ],
    },
    {
        "num": 14,
        "title": "Integrações — Conectar Ferramentas Externas",
        "category": "Integrações",
        "level": "Intermédio",
        "duration": "4 min",
        "objective": "Configurar integrações com Microsoft 365, Jira, Slack e outras ferramentas para automatizar a recolha de evidências.",
        "topics": ["Integration Hub", "Conectar Microsoft 365", "Conectar Jira", "Testar e monitorizar integração"],
        "script": [
            ("Cena 1 — Integration Hub (0:00–0:30)", [
                "Narrador: O Integration Hub do iComply permite conectar as ferramentas que já usa para recolher evidências automaticamente.",
                "Ação: Aceder a Integration Hub. Mostrar as integrações disponíveis: Microsoft 365, Google Workspace, AWS, Jira, Slack, GitHub, etc.",
            ]),
            ("Cena 2 — Conectar Microsoft 365 (0:30–2:00)", [
                "Ação: Clicar em 'Microsoft 365'. Clicar 'Conectar'. O sistema redirige para o ecrã de autorização da Microsoft.",
                "Narrador: Após autorização, o iComply tem acesso de leitura para recolher evidências: políticas de acesso condicional, configuração de MFA, etc.",
                "Ação: Voltar ao iComply. A integração aparece como 'Conectada'. Clicar 'Sincronizar' para importar evidências.",
            ]),
            ("Cena 3 — Configurar Jira (2:00–3:15)", [
                "Ação: Clicar Jira. Inserir URL da instância e API Key.",
                "Narrador: Com o Jira conectado, os tickets de segurança fechados são importados automaticamente como evidências.",
            ]),
            ("Cena 4 — Monitorizar integrações (3:15–4:00)", [
                "Narrador: O log de integrações mostra quando foi a última sincronização, quantas evidências foram importadas e eventuais erros.",
                "Ação: Mostrar o log de sincronização.",
            ]),
        ],
    },
    {
        "num": 15,
        "title": "DORA — Resiliência Operacional Digital",
        "category": "DORA",
        "level": "Avançado",
        "duration": "6 min",
        "objective": "Implementar os requisitos do Digital Operational Resilience Act (DORA) para entidades financeiras.",
        "topics": ["Âmbito do DORA", "Gestão de risco TIC", "Testes de resiliência", "Relatórios de incidentes TIC"],
        "script": [
            ("Cena 1 — O que é o DORA (0:00–1:00)", [
                "Narrador: O DORA (Regulamento EU 2022/2554) aplica-se a entidades financeiras e impõe requisitos de resiliência operacional digital a partir de Janeiro de 2025.",
                "Narrador: O iComply inclui o framework DORA completo com todos os artigos e requisitos mapeados.",
            ]),
            ("Cena 2 — Activar DORA (1:00–2:00)", [
                "Ação: Activar o framework DORA. O sistema cria os domínios DORA: Gestão de Risco TIC, Incidentes, Testes, Riscos de Terceiros, Partilha de Informação.",
            ]),
            ("Cena 3 — Gestão de risco TIC (2:00–3:30)", [
                "Narrador: O DORA exige uma estrutura de governação do risco TIC com papéis definidos ao nível do Órgão de Administração.",
                "Ação: Criar activos TIC críticos. Associar riscos. Mapear a controlos DORA.",
            ]),
            ("Cena 4 — Registo e notificação de incidentes (3:30–5:00)", [
                "Narrador: Os incidentes TIC graves devem ser notificados às autoridades em 4 horas (relatório inicial) e 72 horas (relatório intermédio).",
                "Ação: Criar incidente TIC. Preencher campos DORA. Gerar relatório de notificação.",
            ]),
            ("Cena 5 — Relatório DORA (5:00–6:00)", [
                "Narrador: O relatório anual DORA resume o nível de maturidade, os incidentes do ano e o programa de testes de resiliência.",
                "Ação: Gerar relatório DORA anual.",
            ]),
        ],
    },
    {
        "num": 16,
        "title": "Administração — Utilizadores e Permissões",
        "category": "Administração",
        "level": "Intermédio",
        "duration": "4 min",
        "objective": "Gerir utilizadores, papéis e permissões na organização do iComply.",
        "topics": ["Convidar utilizadores", "Papéis e permissões", "Departamentos", "Auditoria de acessos"],
        "script": [
            ("Cena 1 — Gestão de utilizadores (0:00–0:30)", [
                "Narrador: Como administrador, pode convidar utilizadores, definir papéis e gerir permissões.",
                "Ação: Aceder a Administração → Utilizadores.",
            ]),
            ("Cena 2 — Convidar e configurar utilizador (0:30–2:00)", [
                "Ação: Clicar 'Convidar Utilizador'. Inserir email. Selecionar papel: Administrador, Gestor de Conformidade, Colaborador, Auditor.",
                "Narrador: O papel determina o que o utilizador pode ver e fazer. Um Colaborador pode ver as suas tarefas mas não gerir riscos.",
                "Ação: O utilizador recebe email de convite. Mostrar o processo de activação.",
            ]),
            ("Cena 3 — Departamentos e estrutura (2:00–3:00)", [
                "Narrador: Organize os utilizadores por departamentos para facilitar a atribuição de tarefas e responsabilidades.",
                "Ação: Criar departamento 'IT', 'Jurídico', 'Operações'. Atribuir utilizadores.",
            ]),
            ("Cena 4 — Auditoria de acessos (3:00–4:00)", [
                "Narrador: O log de actividade regista todas as acções dos utilizadores — fundamental para conformidade e investigação de incidentes.",
                "Ação: Mostrar o log de actividade com filtros por utilizador, data e acção.",
            ]),
        ],
    },
    {
        "num": 17,
        "title": "Inteligência Regulatória — Monitorizar Alterações Legais",
        "category": "Inteligência Regulatória",
        "level": "Avançado",
        "duration": "4 min",
        "objective": "Usar o módulo de Inteligência Regulatória para monitorizar alterações legislativas e regulamentares relevantes.",
        "topics": ["Feed regulatório", "Alertas personalizados", "Impacto na conformidade", "Histórico regulatório"],
        "script": [
            ("Cena 1 — Inteligência Regulatória (0:00–0:45)", [
                "Narrador: O módulo de Inteligência Regulatória monitorizará automaticamente publicações de reguladores europeus, nacionais e sectoriais.",
                "Ação: Aceder ao módulo. Mostrar o feed de actualizações regulatórias recentes.",
            ]),
            ("Cena 2 — Configurar alertas (0:45–2:00)", [
                "Ação: Configurar alertas: selecionar reguladores (CNPD, Banco de Portugal, CNCS, ENISA), temas (RGPD, NIS2, DORA, PSD3).",
                "Narrador: Quando sai uma nova publicação relevante, recebe uma notificação no iComply e por email.",
            ]),
            ("Cena 3 — Avaliar impacto (2:00–3:15)", [
                "Narrador: Para cada actualização regulatória, o iComply ajuda-o a avaliar o impacto na sua conformidade.",
                "Ação: Abrir uma actualização. Usar o assistente para mapear o impacto a controlos e políticas existentes.",
                "Narrador: Se a actualização exigir mudanças, crie tarefas directamente a partir daqui.",
            ]),
            ("Cena 4 — Histórico e relatório (3:15–4:00)", [
                "Narrador: O histórico regulatório mostra todas as actualizações que afectaram a sua organização — útil para auditorias.",
            ]),
        ],
    },
    {
        "num": 18,
        "title": "Relatórios de Conformidade — Criar e Exportar",
        "category": "Relatórios",
        "level": "Intermédio",
        "duration": "4 min",
        "objective": "Gerar relatórios de conformidade detalhados para diferentes audiências: gestão, auditores e reguladores.",
        "topics": ["Tipos de relatório", "Configurar relatório", "Exportar PDF/Word", "Agendar relatórios automáticos"],
        "script": [
            ("Cena 1 — Tipos de relatório (0:00–0:45)", [
                "Narrador: O iComply oferece vários tipos de relatório: relatório de conformidade geral, por framework, de riscos, de auditorias e relatório executivo.",
                "Ação: Mostrar o menu de relatórios.",
            ]),
            ("Cena 2 — Configurar e gerar relatório (0:45–2:15)", [
                "Ação: Criar relatório de conformidade ISO 27001. Selecionar período: Trimestre actual. Selecionar secções a incluir.",
                "Narrador: Pode personalizar o logotipo, cabeçalho e as secções incluídas no relatório.",
                "Ação: Clicar 'Gerar'. O relatório aparece em segundos.",
            ]),
            ("Cena 3 — Exportar e partilhar (2:15–3:15)", [
                "Ação: Exportar em PDF. Partilhar por link seguro com um auditor externo.",
                "Narrador: Os links de partilha expiram automaticamente após o período configurado.",
            ]),
            ("Cena 4 — Relatórios automáticos (3:15–4:00)", [
                "Narrador: Configure relatórios automáticos mensais ou trimestrais enviados por email aos destinatários definidos.",
                "Ação: Criar agendamento de relatório mensal de conformidade.",
            ]),
        ],
    },
    {
        "num": 19,
        "title": "Gestão de Fornecedores e Cadeia de Fornecimento",
        "category": "Fornecedores",
        "level": "Avançado",
        "duration": "5 min",
        "objective": "Avaliar e monitorizar o risco de terceiros e fornecedores críticos conforme exigido pela NIS2 e DORA.",
        "topics": ["Registo de fornecedores", "Avaliação de risco de terceiros", "Questionários de due diligence", "Contratos e SLAs"],
        "script": [
            ("Cena 1 — Risco de terceiros (0:00–0:45)", [
                "Narrador: A NIS2 e o DORA exigem que as organizações avaliem e monitorizem o risco associado a fornecedores e prestadores de serviços TIC.",
            ]),
            ("Cena 2 — Registo de fornecedores (0:45–2:00)", [
                "Ação: Aceder a Fornecedores. Criar fornecedor: nome, categoria, serviços prestados, criticidade (Crítico/Alto/Médio/Baixo).",
                "Narrador: Classifique os fornecedores por criticidade para priorizar a avaliação.",
            ]),
            ("Cena 3 — Questionário de due diligence (2:00–3:30)", [
                "Ação: Enviar questionário de segurança ao fornecedor. O fornecedor recebe link e preenche no portal.",
                "Narrador: O questionário cobre: certificações de segurança, políticas de dados, controlo de acessos, respostas a incidentes.",
                "Ação: Ver as respostas do fornecedor. O sistema calcula automaticamente o score de risco.",
            ]),
            ("Cena 4 — Monitorização contínua (3:30–5:00)", [
                "Narrador: Configure alertas para reavaliar fornecedores críticos anualmente ou quando ocorrer um incidente.",
                "Ação: Mostrar o dashboard de risco de fornecedores com score por fornecedor e alertas activos.",
            ]),
        ],
    },
    {
        "num": 20,
        "title": "Portal de Auditores — Vista Externa",
        "category": "Portal de Auditores",
        "level": "Intermédio",
        "duration": "4 min",
        "objective": "Demonstrar o Portal de Auditores — a vista segura para auditores externos acederem a evidências e controlos.",
        "topics": ["Convidar auditor externo", "Vista do auditor", "Recolher evidências", "Gerar relatório final"],
        "script": [
            ("Cena 1 — Portal de Auditores (0:00–0:30)", [
                "Narrador: O Portal de Auditores do iComply permite convidar auditores externos com acesso limitado e seguro à informação necessária para a auditoria.",
            ]),
            ("Cena 2 — Convidar auditor (0:30–1:30)", [
                "Ação: Em Auditorias, selecionar auditoria. Clicar 'Convidar Auditor'. Inserir email. Definir âmbito de acesso e data de expiração.",
                "Narrador: O auditor recebe um link seguro por email. Não precisa de criar conta — acede directamente.",
            ]),
            ("Cena 3 — Vista do auditor (1:30–3:00)", [
                "Ação: Simular a vista do auditor. Mostra apenas os controlos e evidências do âmbito definido.",
                "Narrador: O auditor pode avaliar cada controlo, adicionar comentários e solicitar evidências adicionais.",
                "Ação: O auditor marca um controlo como 'Não Conforme' e adiciona comentário.",
            ]),
            ("Cena 4 — Comunicação e relatório (3:00–4:00)", [
                "Narrador: O sistema de mensagens permite comunicação directa entre o auditor e a equipa interna dentro da plataforma.",
                "Ação: O auditor envia mensagem a solicitar evidência adicional. A equipa interna responde com o ficheiro.",
                "Narrador: No final, o auditor gera o relatório de auditoria directamente do portal.",
            ]),
        ],
    },
]


def add_heading(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    if color:
        for run in p.runs:
            run.font.color.rgb = color
    return p


def add_styled_paragraph(doc, text, bold=False, color=None, size=None, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    if bold:
        run.font.bold = True
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    return p


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def main():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Title page
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("iComply Academy")
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = BRAND_BLUE

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle_p.add_run("Scripts de Vídeo — Centro de Formação")
    run.font.size = Pt(18)
    run.font.color.rgb = BRAND_LIGHT

    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta_p.add_run("20 Vídeos · Português Europeu · Versão 1.0")
    run.font.size = Pt(11)
    run.font.color.rgb = GRAY_TEXT

    doc.add_paragraph()

    # Summary table
    summary_p = doc.add_paragraph()
    run = summary_p.add_run("Índice de Vídeos")
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = BRAND_BLUE

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(["#", "Título", "Categoria", "Nível"]):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].font.bold = True
        hdr[i].paragraphs[0].runs[0].font.color.rgb = WHITE
        set_cell_bg(hdr[i], "1E3A8A")

    for v in VIDEOS:
        row = table.add_row().cells
        row[0].text = str(v["num"])
        row[1].text = v["title"]
        row[2].text = v["category"]
        row[3].text = v["level"]

    doc.add_page_break()

    # Video scripts
    for v in VIDEOS:
        # Video header
        h = doc.add_heading(f'Vídeo {v["num"]} — {v["title"]}', level=1)
        for run in h.runs:
            run.font.color.rgb = BRAND_BLUE
            run.font.size = Pt(16)

        # Metadata table
        meta_table = doc.add_table(rows=1, cols=4)
        meta_table.style = 'Table Grid'
        for cell, (label, value) in zip(meta_table.rows[0].cells, [
            ("Categoria", v["category"]),
            ("Nível", v["level"]),
            ("Duração", v["duration"]),
            ("Objectivo", ""),
        ]):
            p = cell.paragraphs[0]
            r1 = p.add_run(label + ": ")
            r1.font.bold = True
            r1.font.size = Pt(9)
            r1.font.color.rgb = GRAY_TEXT
            r2 = p.add_run(value)
            r2.font.size = Pt(9)
            set_cell_bg(cell, "EFF6FF")

        # Objective
        obj_p = doc.add_paragraph()
        run = obj_p.add_run("Objectivo: ")
        run.font.bold = True
        run.font.color.rgb = BRAND_BLUE
        run2 = obj_p.add_run(v["objective"])
        run2.font.color.rgb = DARK_TEXT

        # Topics
        topics_p = doc.add_paragraph()
        run = topics_p.add_run("Tópicos: ")
        run.font.bold = True
        run2 = topics_p.add_run(" · ".join(v["topics"]))
        run2.font.color.rgb = GRAY_TEXT
        run2.font.size = Pt(10)

        doc.add_paragraph()

        # Scenes
        for scene_title, lines in v["script"]:
            scene_p = doc.add_paragraph()
            scene_p.paragraph_format.left_indent = Cm(0)
            run = scene_p.add_run(scene_title)
            run.font.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = BRAND_LIGHT

            for line in lines:
                if line.startswith("Ação:") or line.startswith("Acção:"):
                    p = doc.add_paragraph(style='List Bullet')
                    p.paragraph_format.left_indent = Cm(1)
                    r1 = p.add_run("ACÇÃO: ")
                    r1.font.bold = True
                    r1.font.color.rgb = GREEN_TEXT
                    r1.font.size = Pt(10)
                    r2 = p.add_run(line.replace("Ação: ", "").replace("Acção: ", ""))
                    r2.font.size = Pt(10)
                    r2.font.color.rgb = DARK_TEXT
                else:
                    p = doc.add_paragraph(style='List Bullet')
                    p.paragraph_format.left_indent = Cm(1)
                    r1 = p.add_run("NARRADOR: ")
                    r1.font.bold = True
                    r1.font.color.rgb = BRAND_BLUE
                    r1.font.size = Pt(10)
                    r2 = p.add_run(line.replace("Narrador: ", ""))
                    r2.font.size = Pt(10)
                    r2.font.color.rgb = DARK_TEXT

            doc.add_paragraph()

        # Production notes
        notes_p = doc.add_paragraph()
        run = notes_p.add_run("Notas de Produção: ")
        run.font.bold = True
        run.font.color.rgb = GRAY_TEXT
        run.font.size = Pt(9)
        run2 = notes_p.add_run(
            f"Gravar o ecrã do iComply com resolução 1920x1080. "
            f"Narração gravada separadamente em português europeu. "
            f"Duração alvo: {v['duration']}. Legendas SRT a incluir."
        )
        run2.font.size = Pt(9)
        run2.font.color.rgb = GRAY_TEXT

        doc.add_page_break()

    doc.save(OUTPUT)
    print(f"Documento guardado em: {OUTPUT}")


if __name__ == "__main__":
    main()
